#!/usr/bin/env python3
"""Apply deterministic RUDN-style academic formatting to Quarto DOCX output."""

from pathlib import Path
import re
import sys

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


FONT = "Times New Roman"
INK = RGBColor(0x17, 0x25, 0x30)
ACCENT = RGBColor(0x12, 0x5B, 0x74)


def set_run_font(run, size=None, bold=None, italic=None, color=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_after(paragraph, text, style=None):
    new_paragraph = paragraph._parent.add_paragraph(text, style=style)
    paragraph._p.addnext(new_paragraph._p)
    return new_paragraph


def add_page_field(paragraph):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend([begin, instr, separate, value, end])
    set_run_font(run, size=10, color=INK)


def iter_tables(parent):
    for table in parent.tables:
        yield table
        for row in table.rows:
            for cell in row.cells:
                yield from iter_tables(cell)


def set_table_geometry(table, widths_cm):
    widths_dxa = [round(width / 2.54 * 1440) for width in widths_cm]
    total_dxa = sum(widths_dxa)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total_dxa))

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "0")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            cell.width = Cm(widths_cm[index])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))


def insert_page_break_before(table):
    paragraph = OxmlElement("w:p")
    properties = OxmlElement("w:pPr")
    properties.append(OxmlElement("w:pageBreakBefore"))
    paragraph.append(properties)
    table._tbl.addprevious(paragraph)


def format_document(source: Path, destination: Path):
    doc = Document(source)
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Mm(30)
    section.right_margin = Mm(15)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.header_distance = Mm(10)
    section.footer_distance = Mm(12)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(12)
    normal.font.color.rgb = INK
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.first_line_indent = Cm(1.25)

    for style_name in ["Body Text", "First Paragraph", "Compact", "Abstract"]:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = FONT
            style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
            style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
            style.font.size = Pt(12)
            style.font.color.rgb = INK
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            style.paragraph_format.line_spacing = 1.5
            style.paragraph_format.space_after = Pt(6)

    heading_tokens = {
        "Heading 1": (15, 18, 8),
        "Heading 2": (13, 14, 6),
        "Heading 3": (12, 10, 4),
    }
    for name, (size, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = FONT
        style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
        style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ACCENT
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ["Caption", "Image Caption", "Table Caption"]:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = FONT
            style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
            style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
            style.font.size = Pt(10)
            style.font.italic = True
            style.font.color.rgb = INK
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style.paragraph_format.space_after = Pt(9)

    title = doc.paragraphs[0]
    university = title.insert_paragraph_before(
        "Российский университет дружбы народов имени Патриса Лумумбы"
    )
    university.alignment = WD_ALIGN_PARAGRAPH.CENTER
    university.paragraph_format.space_after = Pt(4)
    set_run_font(university.runs[0], size=13, bold=True, color=INK)

    course = title.insert_paragraph_before("Дисциплина «Имитационное моделирование»")
    course.alignment = WD_ALIGN_PARAGRAPH.CENTER
    course.paragraph_format.space_after = Pt(52)
    set_run_font(course.runs[0], size=12, color=INK)

    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(24)
    title.paragraph_format.space_after = Pt(10)
    for run in title.runs:
        set_run_font(run, size=20, bold=True, color=ACCENT)

    subtitle = doc.paragraphs[3]
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(96)
    for run in subtitle.runs:
        set_run_font(run, size=15, italic=True, color=INK)

    author = next(p for p in doc.paragraphs if p.style.name == "Author")
    author.text = "Выполнила: Куокконен Дарина Андреевна"
    author.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    author.paragraph_format.space_after = Pt(4)
    set_run_font(author.runs[0], size=12, color=INK)

    group = add_after(author, "Группа: НФИбд-01-23")
    group.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    group.paragraph_format.space_after = Pt(70)
    set_run_font(group.runs[0], size=12, color=INK)

    date = next(p for p in doc.paragraphs if p.style.name == "Date")
    date.text = "Москва, 2026"
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date.paragraph_format.space_after = Pt(0)
    set_run_font(date.runs[0], size=12, color=INK)
    date.add_run().add_break(WD_BREAK.PAGE)

    if "Source Code" in doc.styles:
        code = doc.styles["Source Code"]
        code.font.name = "DejaVu Sans Mono"
        code._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "DejaVu Sans Mono")
        code._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "DejaVu Sans Mono")
        code.font.size = Pt(9)
        code.paragraph_format.line_spacing = 1.0
        code.paragraph_format.space_after = Pt(8)
        code.paragraph_format.first_line_indent = Cm(0)
        code.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.widow_control = True
        for run in paragraph.runs:
            if paragraph.style.name != "Source Code":
                set_run_font(run)

    for table in iter_tables(doc):
        if len(table.columns) == 1 and len(table.rows) == 1:
            row_pr = table.rows[0]._tr.get_or_add_trPr()
            if row_pr.find(qn("w:cantSplit")) is None:
                row_pr.append(OxmlElement("w:cantSplit"))
            figure_match = re.search(r"Рисунок\s+(\d+):", table.cell(0, 0).text)
            if figure_match and int(figure_match.group(1)) in {
                4,
                6,
                9,
                14,
                15,
                17,
                19,
                21,
                23,
                25,
                27,
                29,
            }:
                insert_page_break_before(table)
        for row_index, row in enumerate(table.rows):
            for column_index, cell in enumerate(row.cells):
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(2)
                    paragraph.paragraph_format.space_after = Pt(2)
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.paragraph_format.first_line_indent = Cm(0)
                    if len(table.columns) > 1:
                        paragraph.alignment = (
                            WD_ALIGN_PARAGRAPH.LEFT
                            if column_index in (0, len(table.columns) - 1)
                            else WD_ALIGN_PARAGRAPH.CENTER
                        )
                    for run in paragraph.runs:
                        set_run_font(run, size=10, bold=(row_index == 0))

        if len(table.columns) == 3 and len(table.rows) == 11:
            set_table_geometry(table, [4.8, 2.4, 8.5])
        elif len(table.columns) == 3 and len(table.rows) == 6:
            set_table_geometry(table, [5.2, 5.2, 5.3])

    footer = section.footer
    footer.is_linked_to_previous = False
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.clear()
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(footer_paragraph)
    first_footer = section.first_page_footer
    first_footer.paragraphs[0].clear()

    doc.core_properties.title = "Лабораторная работа № 1 — Экспоненциальный рост"
    doc.core_properties.author = "Куокконен Дарина Андреевна"
    doc.core_properties.subject = "Имитационное моделирование, группа НФИбд-01-23"
    doc.core_properties.keywords = "Julia; DifferentialEquations; DrWatson; Jupyter; Quarto"

    destination.parent.mkdir(parents=True, exist_ok=True)
    doc.save(destination)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("usage: polish_docx.py INPUT.docx OUTPUT.docx")
    format_document(Path(sys.argv[1]), Path(sys.argv[2]))
